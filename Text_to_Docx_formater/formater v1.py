from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import yaml

# --- Add hyperlink to a paragraph ---
def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Add a hyperlink to a Word document paragraph.

    Args:
        paragraph: The paragraph object to append the hyperlink to.
        url (str): The URL the hyperlink should point to.
        text (str): The text to display for the hyperlink.
        color (str): The hex color code for the hyperlink text.
        underline (bool): Whether to underline the hyperlink text.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Set color
    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), color)
    rPr.append(color_elem)

    # Set underline if required
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# --- Load YAML rules ---
def load_formatting_rules(yaml_path):
    """
    Load formatting rules from a YAML file.

    Args:
        yaml_path (str): Path to the YAML file.

    Returns:
        dict: Parsed YAML rules.
    """
    with open(yaml_path, "r", encoding="utf-8") as file:
        return yaml.safe_load(file)

# --- Merge styles ---
def merge_styles(parent_style, child_style):
    """
    Merge two style dictionaries, with child overriding parent.

    Args:
        parent_style (dict): Base style.
        child_style (dict): Style to override base.

    Returns:
        dict: Merged style dictionary.
    """
    merged = parent_style.copy()
    merged.update(child_style)
    return merged

# --- Map alignment strings ---
def get_alignment(value):
    """
    Map string alignment to docx constants.

    Args:
        value (str): Alignment as string.

    Returns:
        WD_PARAGRAPH_ALIGNMENT: Corresponding docx alignment.
    """
    return {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    }.get(value, WD_PARAGRAPH_ALIGNMENT.LEFT)

# --- Apply run style ---
def apply_run_style(run, style, text):
    """
    Apply styling attributes to a run.

    Args:
        run: The run object to style.
        style (dict): Style attributes.
        text (str): Text to format.

    Returns:
        str: Transformed text based on style.
    """
    font = run.font
    if style.get("bold"): run.bold = True
    if style.get("italic"): run.italic = True
    if style.get("underline"): run.underline = True
    if style.get("all_caps"): run.all_caps = True
    if style.get("font"): font.name = style["font"]
    if style.get("size"): font.size = Pt(style["size"])
    if style.get("color"):
        font.color.rgb = RGBColor.from_string(style["color"].lstrip("#").upper())
    return text.upper() if style.get("all_caps") else text

# --- Parse tagged text ---
def parse_tagged_text(text, stack, line_number):
    """
    Parse a string with XML-like tags into styled segments.

    Args:
        text (str): Input string to parse.
        stack (list): Tag stack used for nested tags.
        line_number (int): Current line number (for debugging).

    Returns:
        tuple: List of parsed segments and updated tag stack.
    """
    parsed = []

    # Check for <link> self-closing tag
    link_match = re.match(r'<link\s+text="(.+?)"\s+url="(.+?)"\s*/>', text)
    if link_match:
        link_text, link_url = link_match.groups()
        parsed.append((("LINK", link_text, link_url), tuple(stack)))
        return parsed, stack

    # For splitting an annotated string into tags and plain text
    token_pattern = re.compile(r'(</?[^<>]+>)|([^<>]+)')
    buffer = ""

    # For each tag or text segment found
    # The regex captures both opening/closing tags and text content
    for match in token_pattern.finditer(text):
        # Extract the tag or text content
        # match.group(0) is either a tag or text
        # match.group(1) is the tag, match.group(2) is the text
        tag, content = match.groups()
        if tag:
            if buffer:
                parsed.append((buffer, tuple(stack)))
                buffer = ""

            tag_name = tag.strip("</>").strip()
            full_tag = f"<{tag_name}>"
            if tag.startswith('</'):
                if stack and stack[-1] == full_tag:
                    stack.pop()
                else:
                    print(f"⚠️ Warning: Unmatched closing tag </{tag_name}> at line {line_number}: {text}")
            else:
                stack.append(full_tag)
        elif content:
            buffer += content

    if buffer:
        parsed.append((buffer, tuple(stack)))
    return parsed, stack

# --- Format document ---
def format_document(xml_path, yaml_path, output_path):
    """
    Format a Word document based on an XML-style input and YAML rules.

    Args:
        xml_path (str): Path to the input .xml-style file.
        yaml_path (str): Path to the YAML rules file.
        output_path (str): Path where the output .docx will be saved.
    """
    # Load formatting rules from YAML
    rules = load_formatting_rules(yaml_path)
    doc = Document()

    # Read the XML-style input file
    with open(xml_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    stack = []
    # Parse each line of the XML-style input
    for line_number, line in enumerate(lines, start=1):

        # Parse the line to extract segments and their associated tags
        # The stack keeps track of the current tag context (stack is a list of parent tags)
        # Each segment is a tuple of (content, tag_path) (tag_path is a tuple of tags that are open)
        parsed_segments, stack = parse_tagged_text(line.strip(), stack, line_number)
        if not parsed_segments:
            continue

        # Group inline segments into paragraphs based on block-level tags
        block_groups = []
        current_block, current_tag = [], None

        # Iterate through parsed segments to group them by block-level tags
        for content, tag_path in parsed_segments:
            # Determine the block-level tag
            block_tag = next((tag for tag in reversed(tag_path) if rules.get(tag, {}).get("kind") in ["paragraph", "heading"]), "<paragraph>")
            if block_tag != current_tag:
                if current_block:
                    block_groups.append((current_tag, current_block))
                current_block, current_tag = [(content, tag_path)], block_tag
            else:
                current_block.append((content, tag_path))

        if current_block:
            # Append the last block group. Block groups are tuples of (tag, list of segments)
            # Each segment is a tuple of (content, tag_path)
            # The last block group may not be added if it is empty
            block_groups.append((current_tag, current_block))

        # Render each block group
        for block_tag, group in block_groups:
            block_style = rules.get(block_tag, {})

            # Render hyperlink block separately
            if any(isinstance(content, tuple) and content[0] == "LINK" for content, _ in group):
                for content, _ in group:
                    if isinstance(content, tuple) and content[0] == "LINK":
                        para = doc.add_paragraph()
                        add_hyperlink(para, content[2], content[1])
                continue

            # Create a new paragraph for the block
            para = doc.add_paragraph()
            para_format = para.paragraph_format

            # Apply paragraph-level formatting
            if "alignment" in block_style:
                para.alignment = get_alignment(block_style["alignment"])
            if "space_before" in block_style:
                para_format.space_before = Pt(block_style["space_before"])
            if "space_after" in block_style:
                para_format.space_after = Pt(block_style["space_after"])
            if block_style.get("indent") == "indent":
                para_format.left_indent = Pt(18)
            if block_style.get("bullet") == "bullet":
                para.style = 'List Bullet'
            if "line_spacing" in block_style:
                para_format.line_spacing = Pt(block_style["line_spacing"])
            if block_style.get("kind") == "heading":
                para.style = "Heading 1"

            # Render each run of text with appropriate styles
            for i, (text, tag_path) in enumerate(group):
                # Merge styles from parent tags
                effective_style = {}
                for tag in tag_path:
                    effective_style = merge_styles(effective_style, rules.get(tag, {}))
                # Apply run-level formatting
                run = para.add_run()
                styled_text = apply_run_style(run, effective_style, text.strip())
                run.text = (" " + styled_text) if i > 0 else styled_text

    doc.save(output_path)
    print(f"✅ Document saved to {output_path}")

# --- Main ---
if __name__ == "__main__":
    format_document("input.xml", "rules.yaml", "output.docx")
